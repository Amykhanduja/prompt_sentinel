from docx import Document

from connectors.extraction_result import (
    ExtractionResult,
    ExtractedContent,
)
from context.source import ScanSource


def parse_docx(file_path: str) -> ExtractionResult:

    document = Document(file_path)

    items = []

    # Body text
    body_text = []

    for paragraph in document.paragraphs:

        if paragraph.text.strip():
            body_text.append(
                paragraph.text
            )

    if body_text:

        items.append(
            ExtractedContent(
                content="\n".join(body_text),
                source=ScanSource.DOCX
            )
        )


    # Headers

    header_text = []

    for section in document.sections:

        for paragraph in section.header.paragraphs:

            if paragraph.text.strip():

                header_text.append(
                    paragraph.text
                )


    if header_text:

        items.append(
            ExtractedContent(
                content="\n".join(header_text),
                source=ScanSource.DOCX_HEADER
            )
        )


    # Footers

    footer_text = []

    for section in document.sections:

        for paragraph in section.footer.paragraphs:

            if paragraph.text.strip():

                footer_text.append(
                    paragraph.text
                )


    if footer_text:

        items.append(
            ExtractedContent(
                content="\n".join(footer_text),
                source=ScanSource.DOCX_FOOTER
            )
        )

    # Core properties metadata

    properties = document.core_properties

    metadata = []

    if properties.title:
        metadata.append(
            f"Title: {properties.title}"
        )

    if properties.subject:
        metadata.append(
            f"Subject: {properties.subject}"
        )

    if properties.author:
        metadata.append(
            f"Author: {properties.author}"
        )

    if properties.keywords:
        metadata.append(
            f"Keywords: {properties.keywords}"
        )

    if properties.comments:
        metadata.append(
            f"Comments: {properties.comments}"
        )


    if metadata:

        items.append(
            ExtractedContent(
                content="\n".join(metadata),
                source=ScanSource.DOCX_METADATA
            )
        )


    # Tables

    table_text = []

    for table in document.tables:

        for row in table.rows:

            cells = []

            for cell in row.cells:

                if cell.text.strip():
                    cells.append(
                        cell.text
                    )

            if cells:
                table_text.append(
                    " | ".join(cells)
                )


    if table_text:

        items.append(
            ExtractedContent(
                content="\n".join(table_text),
                source=ScanSource.DOCX_TABLE
            )
        )


    return ExtractionResult(
        items=items
    )
